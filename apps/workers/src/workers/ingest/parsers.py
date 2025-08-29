"""
Document and media parsers for the ingest pipeline.
"""

import io
import logging
from typing import Dict, Any, Optional
from pathlib import Path

# Document parsing libraries
try:
    import PyPDF2
    from docx import Document as DocxDocument
    from bs4 import BeautifulSoup
    import markdown
except ImportError as e:
    logging.warning(f"Some parsing libraries not available: {e}")

# Media processing libraries
try:
    from PIL import Image, ExifTags
    import mutagen
except ImportError as e:
    logging.warning(f"Some media libraries not available: {e}")

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for various document formats."""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._parse_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._parse_docx,
            'text/plain': self._parse_text,
            'text/html': self._parse_html,
            'text/markdown': self._parse_markdown,
            'application/json': self._parse_json,
        }
    
    async def parse(self, content: bytes, content_type: str, filename: str) -> Dict[str, Any]:
        """Parse document content based on type."""
        try:
            # Normalize content type
            content_type = content_type.lower() if content_type else ""
            
            # Try to determine type from filename if content_type is generic
            if content_type in ['application/octet-stream', 'binary/octet-stream'] or not content_type:
                content_type = self._guess_type_from_filename(filename)
            
            # Get appropriate parser
            parser_func = self.supported_types.get(content_type)
            
            if not parser_func:
                logger.warning(f"No parser for content type: {content_type}")
                return {
                    "text": "",
                    "metadata": {"error": f"Unsupported content type: {content_type}"},
                    "needs_ocr": self._is_image_type(content_type)
                }
            
            # Parse content
            result = await parser_func(content, filename)
            result["content_type"] = content_type
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {str(e)}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "needs_ocr": False
            }
    
    def _guess_type_from_filename(self, filename: str) -> str:
        """Guess content type from filename extension."""
        suffix = Path(filename).suffix.lower()
        
        type_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.html': 'text/html',
            '.htm': 'text/html',
            '.md': 'text/markdown',
            '.json': 'application/json',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.webp': 'image/webp',
        }
        
        return type_map.get(suffix, 'application/octet-stream')
    
    def _is_image_type(self, content_type: str) -> bool:
        """Check if content type is an image that might need OCR."""
        return content_type.startswith('image/')
    
    async def _parse_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse PDF document."""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            metadata = {
                "pages": len(pdf_reader.pages),
                "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                "author": pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                "subject": pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else '',
            }
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
            
            full_text = "\n\n".join(text_content)
            
            # If no text extracted, might need OCR
            needs_ocr = len(full_text.strip()) < 100
            
            return {
                "text": full_text,
                "metadata": metadata,
                "needs_ocr": needs_ocr
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF {filename}: {str(e)}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "needs_ocr": True  # Fallback to OCR
            }
    
    async def _parse_docx(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse DOCX document."""
        try:
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n[Tables]\n" + "\n".join(tables_text)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables)
            }
            
            return {
                "text": full_text,
                "metadata": metadata,
                "needs_ocr": False
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {filename}: {str(e)}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "needs_ocr": False
            }
    
    async def _parse_text(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse plain text document."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                # Fallback: decode with errors='replace'
                text = content.decode('utf-8', errors='replace')
                encoding_used = 'utf-8 (with errors)'
            
            metadata = {
                "encoding": encoding_used,
                "lines": len(text.splitlines()),
                "characters": len(text)
            }
            
            return {
                "text": text,
                "metadata": metadata,
                "needs_ocr": False
            }
            
        except Exception as e:
            logger.error(f"Error parsing text {filename}: {str(e)}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "needs_ocr": False
            }
    
    async def _parse_html(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse HTML document."""
        try:
            # Decode content
            html_content = content.decode('utf-8', errors='replace')
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract metadata
            title = soup.find('title')
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
            
            metadata = {
                "title": title.get_text() if title else "",
                "description": meta_desc.get('content', '') if meta_desc else "",
                "keywords": meta_keywords.get('content', '') if meta_keywords else "",
            }
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return {
                "text": text,
                "metadata": metadata,
                "needs_ocr": False
            }
            
        except Exception as e:
            logger.error(f"Error parsing HTML {filename}: {str(e)}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "needs_ocr": False
            }
    
    async def _parse_markdown(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse Markdown document."""
        try:
            # Decode content
            md_content = content.decode('utf-8', errors='replace')
            
            # Convert to HTML first to extract metadata
            md = markdown.Markdown(extensions=['meta'])
            html = md.convert(md_content)
            
            # Extract metadata from markdown meta extension
            metadata = {}
            if hasattr(md, 'Meta'):
                for key, value in md.Meta.items():
                    metadata[key] = value[0] if isinstance(value, list) and len(value) == 1 else value
            
            # Use BeautifulSoup to extract clean text from HTML
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)
            
            return {
                "text": text,
                "metadata": metadata,
                "needs_ocr": False
            }
            
        except Exception as e:
            logger.error(f"Error parsing Markdown {filename}: {str(e)}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "needs_ocr": False
            }
    
    async def _parse_json(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse JSON document."""
        try:
            import json
            
            # Decode and parse JSON
            json_content = content.decode('utf-8', errors='replace')
            data = json.loads(json_content)
            
            # Convert JSON to readable text
            text = json.dumps(data, indent=2, ensure_ascii=False)
            
            metadata = {
                "type": type(data).__name__,
                "size": len(str(data)) if isinstance(data, (dict, list)) else 1
            }
            
            return {
                "text": text,
                "metadata": metadata,
                "needs_ocr": False
            }
            
        except Exception as e:
            logger.error(f"Error parsing JSON {filename}: {str(e)}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "needs_ocr": False
            }


class MediaParser:
    """Parser for media files (images, audio, video)."""
    
    def __init__(self):
        self.supported_types = {
            'image/jpeg': self._parse_image,
            'image/png': self._parse_image,
            'image/gif': self._parse_image,
            'image/webp': self._parse_image,
            'audio/mpeg': self._parse_audio,
            'audio/wav': self._parse_audio,
            'audio/mp4': self._parse_audio,
            'video/mp4': self._parse_video,
            'video/webm': self._parse_video,
        }
    
    async def parse(self, content: bytes, content_type: str, filename: str) -> Dict[str, Any]:
        """Parse media content based on type."""
        try:
            content_type = content_type.lower() if content_type else ""
            
            parser_func = self.supported_types.get(content_type)
            
            if not parser_func:
                logger.warning(f"No media parser for content type: {content_type}")
                return {
                    "metadata": {"error": f"Unsupported media type: {content_type}"}
                }
            
            result = await parser_func(content, filename)
            result["content_type"] = content_type
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing media {filename}: {str(e)}")
            return {
                "metadata": {"error": str(e)}
            }
    
    async def _parse_image(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse image file and extract metadata."""
        try:
            image = Image.open(io.BytesIO(content))
            
            metadata = {
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
                "width": image.width,
                "height": image.height,
            }
            
            # Extract EXIF data if available
            if hasattr(image, '_getexif') and image._getexif():
                exif = image._getexif()
                if exif:
                    for tag_id, value in exif.items():
                        tag = ExifTags.TAGS.get(tag_id, tag_id)
                        metadata[f"exif_{tag}"] = str(value)
            
            return {
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error parsing image {filename}: {str(e)}")
            return {
                "metadata": {"error": str(e)}
            }
    
    async def _parse_audio(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse audio file and extract metadata."""
        try:
            # Save to temporary file for mutagen
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name
            
            try:
                audio_file = mutagen.File(temp_path)
                
                metadata = {}
                if audio_file:
                    # Extract common metadata
                    metadata.update({
                        "length": getattr(audio_file.info, 'length', 0),
                        "bitrate": getattr(audio_file.info, 'bitrate', 0),
                        "sample_rate": getattr(audio_file.info, 'sample_rate', 0),
                        "channels": getattr(audio_file.info, 'channels', 0),
                    })
                    
                    # Extract tags
                    if audio_file.tags:
                        for key, value in audio_file.tags.items():
                            metadata[f"tag_{key}"] = str(value[0]) if isinstance(value, list) else str(value)
                
                return {
                    "metadata": metadata
                }
                
            finally:
                os.unlink(temp_path)
                
        except Exception as e:
            logger.error(f"Error parsing audio {filename}: {str(e)}")
            return {
                "metadata": {"error": str(e)}
            }
    
    async def _parse_video(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse video file and extract metadata."""
        try:
            # Similar to audio parsing but for video
            # This would use ffprobe or similar tool to extract video metadata
            metadata = {
                "type": "video",
                "size": len(content)
            }
            
            return {
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error parsing video {filename}: {str(e)}")
            return {
                "metadata": {"error": str(e)}
            }
